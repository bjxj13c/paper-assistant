"""
飞书集成模块 - 文档保存、机器人通知
通过 lark-cli 命令行工具和飞书 Open API 实现。
"""

import json
import subprocess
import tempfile
from pathlib import Path
from datetime import datetime
from typing import Optional

import config
from .agent import AnalysisResult


def _run_lark(args: list) -> dict:
    """执行 lark-cli 命令并返回 JSON 结果。"""
    cmd = ["lark-cli"] + args
    try:
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
        if result.returncode != 0:
            print(f"[飞书] 命令警告: {result.stderr}")
        if result.stdout.strip():
            return json.loads(result.stdout)
    except FileNotFoundError:
        print("[飞书] lark-cli 未找到，已跳过飞书操作")
    except json.JSONDecodeError:
        pass
    except subprocess.TimeoutExpired:
        print("[飞书] 命令超时")
    return {}


def create_doc(title: str, result: AnalysisResult) -> Optional[str]:
    """
    在飞书中创建文档并写入分析报告。
    返回文档 URL。
    """
    md_content = build_report_markdown(result)

    with tempfile.NamedTemporaryFile(
        mode="w", suffix=".md", delete=False, encoding="utf-8"
    ) as f:
        f.write(md_content)
        tmp_path = f.name

    try:
        # 尝试用 lark-cli 创建文档
        api_result = _run_lark([
            "docx", "create",
            "--title", title,
            "--content", tmp_path,
        ])
        if api_result:
            doc_url = api_result.get("url", api_result.get("data", {}).get("url", ""))
            if doc_url:
                print(f"[飞书] 文档已创建: {doc_url}")
                return doc_url

        # 备用方案：飞书文档 API
        result2 = _create_doc_via_api(title, build_report_markdown(result))
        return result2

    finally:
        Path(tmp_path).unlink(missing_ok=True)


def _create_doc_via_api(title: str, content: str) -> Optional[str]:
    """通过飞书 Open API 创建文档。"""
    # 获取 tenant_access_token
    token_result = _run_lark([
        "auth", "tenant-access-token",
    ])
    token = token_result.get("tenant_access_token", "")

    if not token and config.FEISHU_APP_ID:
        # 尝试手动获取 token
        import requests
        resp = requests.post(
            "https://open.feishu.cn/open-apis/auth/v3/tenant_access_token/internal",
            json={"app_id": config.FEISHU_APP_ID, "app_secret": config.FEISHU_APP_SECRET},
        )
        if resp.status_code == 200:
            token = resp.json().get("tenant_access_token", "")

    if not token:
        print("[飞书] 无法获取 access token，请检查飞书配置")
        return None

    # 创建 docx 文档
    import requests
    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/json",
    }
    create_resp = requests.post(
        "https://open.feishu.cn/open-apis/docx/v1/documents",
        headers=headers,
        json={"title": title},
    )
    if create_resp.status_code == 200:
        doc_data = create_resp.json().get("data", {}).get("document", {})
        doc_id = doc_data.get("document_id", "")
        doc_url = f"https://bytedance.feishu.cn/docx/{doc_id}"

        # 写入内容（简化版：写入纯文本）
        if doc_id:
            _write_doc_content(token, doc_id, content)
            print(f"[飞书] 文档已创建: {doc_url}")
            return doc_url

    print(f"[飞书] 创建文档失败: {create_resp.text}")
    return None


def _write_doc_content(token: str, doc_id: str, content: str):
    """向飞书文档写入内容块。"""
    import requests

    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/json",
    }

    # 将内容转为飞书文档块结构（简化版）
    blocks = []
    for line in content.split("\n"):
        line = line.strip()
        if not line:
            blocks.append({
                "block_type": 2,  # 文本块
                "text": {"elements": [{"text_run": {"content": ""}}], "style": {}},
            })
        elif line.startswith("# "):
            blocks.append({
                "block_type": 3,  # H1
                "heading1": {"elements": [{"text_run": {"content": line[2:]}}], "style": {}},
            })
        elif line.startswith("## "):
            blocks.append({
                "block_type": 4,  # H2
                "heading2": {"elements": [{"text_run": {"content": line[3:]}}], "style": {}},
            })
        elif line.startswith("- "):
            blocks.append({
                "block_type": 16,  # 无序列表
                "bullet": {"elements": [{"text_run": {"content": line[2:]}}], "style": {}},
            })
        else:
            blocks.append({
                "block_type": 2,
                "text": {"elements": [{"text_run": {"content": line}}], "style": {}},
            })

    # 飞书写入内容需要分批
    batch_size = 50
    for i in range(0, len(blocks), batch_size):
        batch = blocks[i:i + batch_size]
        try:
            requests.post(
                f"https://open.feishu.cn/open-apis/docx/v1/documents/{doc_id}/blocks/{doc_id}/children",
                headers=headers,
                json={"children": batch, "index": -1},
                timeout=30,
            )
        except Exception as e:
            print(f"[飞书] 写入块失败: {e}")


def send_bot_notification(title: str, summary: str, doc_url: str = "") -> bool:
    """
    通过飞书机器人发送通知（使用 Webhook 方式）。
    """
    webhook = config.FEISHU_BOT_WEBHOOK
    if not webhook:
        print("[飞书] 未配置 FEISHU_BOT_WEBHOOK，跳过通知")
        return False

    import requests

    # 构建卡片元素
    elements = []
    content_text = f"**{title}**"
    if summary:
        content_text += f"\n{summary[:300]}..."
    elements.append({"tag": "div", "text": {"tag": "lark_md", "content": content_text}})

    if doc_url:
        elements.append({"tag": "div", "text": {"tag": "lark_md",
            "content": f"[查看完整报告]({doc_url})"}})
    else:
        elements.append({"tag": "div", "text": {"tag": "lark_md",
            "content": "分析完成"}})

    payload = {
        "msg_type": "interactive",
        "card": {
            "header": {
                "title": {"tag": "plain_text", "content": "论文分析报告"},
                "template": "blue",
            },
            "elements": elements,
        },
    }

    try:
        resp = requests.post(webhook, json=payload, timeout=30)
        if resp.status_code == 200:
            print("[飞书] 通知已发送")
            return True
    except Exception as e:
        print(f"[飞书] 通知发送失败: {e}")
    return False


def save_to_local(result: AnalysisResult, output_dir: str = "") -> str:
    """将分析结果保存为本地 Markdown 和 Word 文件。"""
    output_path = Path(output_dir or config.OUTPUT_DIR)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_title = result.title.replace("/", "_").replace("\\", "_")[:50] or "analysis"

    # Markdown
    md_path = output_path / f"{timestamp}_{safe_title}.md"
    md_path.write_text(build_report_markdown(result), encoding="utf-8")
    print(f"[本地] Markdown 已保存: {md_path}")

    # Word
    docx_path = save_as_docx(result, str(output_path / f"{timestamp}_{safe_title}.docx"))
    print(f"[本地] Word 已保存: {docx_path}")

    return str(md_path)


def build_report_markdown(result: AnalysisResult) -> str:
    """构建格式化的 Markdown 报告。"""
    title = result.title or "论文分析报告"
    lines = [
        f"# 论文分析报告：{title}",
        "",
        f"**生成时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "",
        "---",
        "",
    ]
    if result.abstract:
        lines.append("## 摘要")
        lines.append("")
        lines.append(result.abstract)
        lines.append("")
    if result.keywords:
        lines.append("## 关键词")
        lines.append("")
        lines.append(" | ".join(f"`{k}`" for k in result.keywords))
        lines.append("")
    if result.research_questions:
        lines.append("## 研究问题")
        lines.append("")
        for i, q in enumerate(result.research_questions, 1):
            lines.append(f"{i}. {q}")
        lines.append("")
    if result.methodology:
        lines.append("## 研究方法")
        lines.append("")
        lines.append(result.methodology)
        lines.append("")
    if result.contributions:
        lines.append("## 主要贡献")
        lines.append("")
        for c in result.contributions:
            lines.append(f"- {c}")
        lines.append("")
    if result.strengths:
        lines.append("## 优点")
        lines.append("")
        for s in result.strengths:
            lines.append(f"- {s}")
        lines.append("")
    if result.limitations:
        lines.append("## 局限性")
        lines.append("")
        for l in result.limitations:
            lines.append(f"- {l}")
        lines.append("")
    if result.reading_notes:
        lines.append("## 阅读建议")
        lines.append("")
        lines.append(result.reading_notes)
        lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("*本报告由论文阅读与摘要生成助手自动生成。*")
    return "\n".join(lines)


def _build_markdown_report(title: str, content: str) -> str:
    """构建 Markdown 格式的报告（保留兼容）。"""
    return content


def save_as_docx(result: AnalysisResult, filepath: str = "") -> str:
    """将分析报告导出为 Word (.docx) 文档，包含自动目录（TOC）。"""
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    doc = Document()

    # ===== 样式配置 =====
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    # 设置标题样式
    for i in range(1, 4):
        heading_style = doc.styles[f"Heading {i}"]
        heading_style.font.name = "Microsoft YaHei"
        if i == 1:
            heading_style.font.size = Pt(18)
            heading_style.font.color.rgb = RGBColor(0x1F, 0x23, 0x29)
        elif i == 2:
            heading_style.font.size = Pt(15)
            heading_style.font.color.rgb = RGBColor(0x33, 0x70, 0xFF)
        else:
            heading_style.font.size = Pt(13)
            heading_style.font.color.rgb = RGBColor(0x1F, 0x23, 0x29)

    # ===== 封面标题 =====
    h = doc.add_heading("论文分析报告", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_para.add_run(f"论文标题：{result.title or '未知'}").font.size = Pt(12)
    doc.add_paragraph()

    time_para = doc.add_paragraph()
    time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    time_para.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").font.size = Pt(10)
    time_para.runs[0].font.color.rgb = RGBColor(0x8F, 0x95, 0x9E)

    # ===== 目录页 =====
    doc.add_paragraph()  # 空行

    # 插入分页符
    run = doc.add_paragraph().add_run()
    run._element.append(parse_xml(f'<w:br {nsdecls("w")} w:type="page"/>'))

    # 目录标题
    toc_heading = doc.add_paragraph()
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_heading.add_run("目  录")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.name = "Microsoft YaHei"
    doc.add_paragraph()

    # 插入 TOC 字段
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fldChar_begin)

    run2 = paragraph.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
    run2._element.append(instrText)

    run3 = paragraph.add_run()
    fldChar_separate = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._element.append(fldChar_separate)

    run4 = paragraph.add_run("（在 Word 中右键此处 → 更新域 → 更新整个目录）")
    run4.font.size = Pt(9)
    run4.font.color.rgb = RGBColor(0x8F, 0x95, 0x9E)

    run5 = paragraph.add_run()
    fldChar_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._element.append(fldChar_end)

    # 分页
    run = doc.add_paragraph().add_run()
    run._element.append(parse_xml(f'<w:br {nsdecls("w")} w:type="page"/>'))

    # ===== 正文内容（使用 Heading 样式，TOC 自动识别） =====

    def add_section(heading, level, body_list):
        """添加章节，使用正确的 Heading 样式以便 TOC 识别。"""
        doc.add_heading(heading, level=level)
        if isinstance(body_list, str):
            doc.add_paragraph(body_list)
        elif isinstance(body_list, list):
            for item in body_list:
                doc.add_paragraph(item, style="List Bullet")

    if result.abstract:
        add_section("摘要", 1, result.abstract)

    if result.keywords:
        doc.add_heading("关键词", level=1)
        p = doc.add_paragraph()
        for i, k in enumerate(result.keywords):
            if i > 0:
                p.add_run("  |  ")
            run = p.add_run(k)
            run.font.color.rgb = RGBColor(0x33, 0x70, 0xFF)
            run.font.size = Pt(11)

    if result.research_questions:
        add_section("研究问题", 1,
            [f"{i}. {q}" for i, q in enumerate(result.research_questions, 1)])

    if result.methodology:
        add_section("研究方法", 1, result.methodology)

    if result.contributions:
        add_section("主要贡献", 1, result.contributions)

    if result.strengths:
        add_section("优点", 1, result.strengths)

    if result.limitations:
        add_section("局限性", 1, result.limitations)

    if result.reading_notes:
        add_section("阅读建议", 1, result.reading_notes)

    # ===== 页脚 =====
    doc.add_paragraph()
    doc.add_paragraph("—" * 40)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("本报告由论文阅读与摘要生成助手自动生成。")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # ===== 保存 =====
    if not filepath:
        output_path = Path(config.OUTPUT_DIR)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_title = result.title.replace("/", "_").replace("\\", "_")[:50] or "analysis"
        filepath = str(output_path / f"{timestamp}_{safe_title}.docx")

    doc.save(filepath)
    return filepath
